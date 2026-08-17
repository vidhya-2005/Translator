import base64
import html
import os
import tempfile
from io import BytesIO

import requests
from docx import Document
from googletrans import LANGUAGES
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

from services.gemini_service import _call, _parse_json, translate_segments

SUPPORTED_EXTENSIONS = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": None,
}

IMAGE_MODEL = "gemini-3.1-flash-image"
FONT_BASE_URL = "https://raw.githubusercontent.com/notofonts/noto-fonts/main/hinted/ttf"
FONT_MAP = {
    "Tamil": "NotoSansTamil",
    "Hindi": "NotoSansDevanagari",
    "Marathi": "NotoSansDevanagari",
    "Nepali": "NotoSansDevanagari",
    "Bengali": "NotoSansBengali",
    "Assamese": "NotoSansBengali",
    "Gujarati": "NotoSansGujarati",
    "Punjabi": "NotoSansGurmukhi",
    "Kannada": "NotoSansKannada",
    "Malayalam": "NotoSansMalayalam",
    "Telugu": "NotoSansTelugu",
    "Odia": "NotoSansOriya",
}


def get_file_type(filename):
    extension = os.path.splitext(filename.lower())[1]
    if extension not in SUPPORTED_EXTENSIONS:
        return None, extension
    return SUPPORTED_EXTENSIONS[extension], extension


def extract_docx_text(path):
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def extract_pdf_text(path):
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def encode_file(path):
    with open(path, "rb") as file:
        return base64.b64encode(file.read()).decode("utf-8")


def _paragraphs(document):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            paragraphs.extend(nested_cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def translate_docx_preserving_format(path, output_path, translate_text_func, target_language, source_language="auto"):
    document = Document(path)
    candidates = []
    for paragraph in _paragraphs(document):
        for run in paragraph.runs:
            if run.text and run.text.strip():
                candidates.append((run, run.text))

    if not candidates:
        raise ValueError("No readable text found in the Word document.")

    translations = translate_segments(
        [text for _, text in candidates],
        target_language,
        source_language,
    )

    if len(translations) != len(candidates):
        raise ValueError("Word translation returned an incomplete result.")

    for (run, _), translated in zip(candidates, translations):
        run.text = str(translated)

    document.save(output_path)


def _visual_prompt(target_language, source_language="auto"):
    source = "Detect the source language automatically." if source_language == "auto" else f"The source language is {LANGUAGES.get(source_language, source_language)}."
    return (
        "Read every visible human-readable text region in this image. "
        f"{source} Translate every region to {target_language}. "
        "Return ONLY JSON with keys detected_language_name, detected_language_code, and regions. "
        "Each region must contain text, translation, x, y, width, height. "
        "Coordinates are normalized 0-1000 relative to image width and height."
    )


def _analyze_visual(path, mime_type, target_language, source_language):
    return _parse_json(_call({"contents": [{"parts": [
        {"text": _visual_prompt(target_language, source_language)},
        {"inlineData": {"mimeType": mime_type, "data": encode_file(path)}}
    ]}]}, {"responseMimeType": "application/json"}))


def _font_cache_dir():
    path = os.path.join(tempfile.gettempdir(), "translator-fonts")
    os.makedirs(path, exist_ok=True)
    return path


def _font(size, target_language="English", bold=False):
    family = FONT_MAP.get(target_language)
    if family:
        filename = f"{family}-{'Bold' if bold else 'Regular'}.ttf"
        cached = os.path.join(_font_cache_dir(), filename)
        if not os.path.exists(cached):
            url = f"{FONT_BASE_URL}/{family}/{filename}"
            try:
                response = requests.get(url, timeout=20)
                response.raise_for_status()
                with open(cached, "wb") as file:
                    file.write(response.content)
            except requests.RequestException:
                pass
        if os.path.exists(cached):
            return ImageFont.truetype(cached, max(10, size))

    for path in (
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ):
        if os.path.exists(path):
            return ImageFont.truetype(path, max(10, size))
    return ImageFont.load_default()


def _fit_font(draw, text, box_width, box_height, target_language):
    size = max(10, int(box_height * 0.72))
    while size >= 10:
        font = _font(size, target_language)
        bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=max(1, int(size * 0.08)))
        if bbox[2] - bbox[0] <= box_width * 0.94 and bbox[3] - bbox[1] <= box_height * 0.90:
            return font
        size -= 2
    return _font(10, target_language)


def _render_translated_image(image, analysis, target_language="English"):
    image = image.convert("RGB")
    draw = ImageDraw.Draw(image)
    width, height = image.size

    for region in analysis.get("regions", []):
        translation = str(region.get("translation", "")).strip()
        if not translation:
            continue

        x = int(float(region.get("x", 0)) / 1000 * width)
        y = int(float(region.get("y", 0)) / 1000 * height)
        w = max(10, int(float(region.get("width", 100)) / 1000 * width))
        h = max(10, int(float(region.get("height", 40)) / 1000 * height))
        x2 = min(width, x + w)
        y2 = min(height, y + h)

        # Sample the surrounding pixels instead of assuming a white background.
        samples = []
        for sx, sy in ((x - 2, y + h // 2), (x2 + 2, y + h // 2), (x + w // 2, y - 2), (x + w // 2, y2 + 2)):
            if 0 <= sx < width and 0 <= sy < height:
                samples.append(image.getpixel((sx, sy)))
        background = tuple(sum(pixel[i] for pixel in samples) // len(samples) for i in range(3)) if samples else (255, 255, 255)

        draw.rectangle((x, y, x2, y2), fill=background)
        font = _fit_font(draw, translation, max(10, x2 - x), max(10, y2 - y), target_language)
        draw.multiline_text(
            (x + 3, y + 2),
            translation,
            fill=(0, 0, 0),
            font=font,
            spacing=max(1, int(font.size * 0.08)),
            align="left",
        )
    return image


def _result(analysis):
    regions = analysis.get("regions", [])
    return {
        "detected_language_name": analysis.get("detected_language_name", "Unknown"),
        "detected_language_code": analysis.get("detected_language_code", ""),
        "transcription": "\n".join(str(r.get("text", "")) for r in regions if r.get("text")),
        "translation": "\n".join(str(r.get("translation", "")) for r in regions if r.get("translation")),
    }


def _aspect_ratio(width, height):
    ratios = {
        "1:1": 1.0,
        "1:4": 0.25,
        "4:1": 4.0,
        "1:8": 0.125,
        "8:1": 8.0,
        "2:3": 2 / 3,
        "3:2": 3 / 2,
        "3:4": 3 / 4,
        "4:3": 4 / 3,
        "4:5": 4 / 5,
        "5:4": 5 / 4,
        "9:16": 9 / 16,
        "16:9": 16 / 9,
        "21:9": 21 / 9,
    }
    value = width / height
    return min(ratios, key=lambda key: abs(ratios[key] - value))


def _gemini_edit_image(path, mime_type, target_language):
    image_data = encode_file(path)
    with Image.open(path) as source:
        width, height = source.size
    prompt = (
        f"Translate ONLY the visible human-readable text in this image into {target_language}. "
        "This is an image-editing task, not a redesign. Preserve the original image as closely as possible. "
        "Keep the exact composition, objects, people, illustrations, background, colors, borders, spacing, "
        "decorations, lighting, perspective and overall layout unchanged. Change ONLY the written text. "
        "Keep each translated text in the same location and approximately the same size, weight, alignment "
        "and visual style as the original. Do not add, remove, invent, summarize or reorder content. "
        "Render the requested script with proper Unicode glyphs; NEVER use square boxes or missing-glyph symbols."
    )
    response = requests.post(
        "https://generativelanguage.googleapis.com/v1beta/interactions",
        headers={"Content-Type": "application/json", "x-goog-api-key": os.environ.get("GEMINI_API_KEY", "")},
        json={
            "model": IMAGE_MODEL,
            "input": [
                {"type": "image", "mime_type": mime_type, "data": image_data},
                {"type": "text", "text": prompt},
            ],
            "response_format": {
                "type": "image",
                "mime_type": "image/png",
                "aspect_ratio": _aspect_ratio(width, height),
                "image_size": "1K",
            },
        },
        timeout=180,
    )
    if response.status_code not in (200, 201):
        try:
            message = response.json().get("error", {}).get("message", response.text)
        except ValueError:
            message = response.text
        raise ValueError(f"Gemini image editing error ({response.status_code}): {message}")

    data = response.json()
    encoded = None
    image = data.get("output_image")
    if image:
        encoded = image.get("data")
    if not encoded:
        for step in data.get("steps", []):
            if step.get("type") != "model_output":
                continue
            for item in step.get("content", []):
                if item.get("type") == "image" and item.get("data"):
                    encoded = item["data"]
                    break
            if encoded:
                break
    if not encoded:
        raise ValueError("Gemini image editing returned no image output.")

    edited = Image.open(BytesIO(base64.b64decode(encoded))).convert("RGB")
    if edited.size != (width, height):
        edited = edited.resize((width, height), Image.Resampling.LANCZOS)
    return edited


def translate_image_file(path, output_path, target_language, source_language, mime_type):
    analysis = _analyze_visual(path, mime_type, target_language, source_language)
    try:
        translated = _gemini_edit_image(path, mime_type, target_language)
    except Exception:
        # Keep a deterministic local fallback. The target-script font is downloaded
        # from the Noto project when it is not already installed on the server.
        translated = _render_translated_image(Image.open(path), analysis, target_language)

    if mime_type == "image/jpeg":
        translated.save(output_path, format="JPEG", quality=95, optimize=True)
    else:
        translated.save(output_path, format="PNG", optimize=True)
    return _result(analysis)


def translate_pdf_file(path, output_path, target_language, source_language):
    import pymupdf

    doc = pymupdf.open(path)
    all_transcription = []
    all_translation = []
    detected_name = "Auto-detected" if source_language == "auto" else LANGUAGES.get(source_language, source_language).capitalize()
    detected_code = "" if source_language == "auto" else source_language

    try:
        for page in doc:
            text_blocks = []
            for block in page.get_text("dict", flags=11).get("blocks", []):
                if block.get("type") != 0:
                    continue
                lines = block.get("lines", [])
                text = "\n".join(
                    span.get("text", "")
                    for line in lines
                    for span in line.get("spans", [])
                    if span.get("text")
                ).strip()
                if not text:
                    continue
                first_span = next((span for line in lines for span in line.get("spans", []) if span.get("text")), {})
                text_blocks.append({
                    "rect": pymupdf.Rect(block["bbox"]),
                    "text": text,
                    "size": float(first_span.get("size", 11)),
                    "color": int(first_span.get("color", 0)),
                    "flags": int(first_span.get("flags", 0)),
                    "dir": tuple(lines[0].get("dir", (1, 0))) if lines else (1, 0),
                })

            if not text_blocks:
                pix = page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5), alpha=False)
                image = Image.open(BytesIO(pix.tobytes("png")))
                temp_path = f"{path}.page-{page.number}.png"
                image.save(temp_path, format="PNG")
                try:
                    analysis = _analyze_visual(temp_path, "image/png", target_language, source_language)
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                translated = _render_translated_image(image, analysis, target_language)
                image_bytes = BytesIO()
                translated.save(image_bytes, format="PNG")
                page.clean_contents()
                page.insert_image(page.rect, stream=image_bytes.getvalue())
                page_result = _result(analysis)
                detected_name = analysis.get("detected_language_name", detected_name)
                detected_code = analysis.get("detected_language_code", detected_code)
                if page_result["transcription"]:
                    all_transcription.append(page_result["transcription"])
                if page_result["translation"]:
                    all_translation.append(page_result["translation"])
                continue

            translations = translate_segments([block["text"] for block in text_blocks], target_language, source_language)
            if len(translations) != len(text_blocks):
                raise ValueError("PDF translation returned an incomplete result.")

            for block in text_blocks:
                page.add_redact_annot(block["rect"] + (-0.5, -0.5, 0.5, 0.5), fill=False, cross_out=False)
            page.apply_redactions(images=0, graphics=0)

            for block, translated in zip(text_blocks, translations):
                translated = str(translated).strip()
                if not translated:
                    continue
                color = block["color"]
                rgb = (((color >> 16) & 255) / 255, ((color >> 8) & 255) / 255, (color & 255) / 255)
                flags = block["flags"]
                weight = "bold" if flags & (1 << 4) else "normal"
                style = "italic" if flags & (1 << 1) else "normal"
                dx, dy = block["dir"]
                angle = 90 if abs(dx) < 0.01 and dy > 0 else 270 if abs(dx) < 0.01 else 180 if dx < 0 else 0
                html_text = html.escape(translated).replace("\n", "<br>")
                html_box = (
                    f'<div style="font-family:sans-serif;font-size:{block["size"]:.2f}pt;'
                    f'font-weight:{weight};font-style:{style};color:rgb('
                    f'{int(rgb[0]*255)},{int(rgb[1]*255)},{int(rgb[2]*255)});">'
                    f'{html_text}</div>'
                )
                page.insert_htmlbox(block["rect"], html_box, rotate=angle, overlay=True)
                all_transcription.append(block["text"])
                all_translation.append(translated)

        doc.save(output_path, garbage=4, deflate=True)
    finally:
        doc.close()

    return {
        "detected_language_name": detected_name,
        "detected_language_code": detected_code,
        "transcription": "\n".join(all_transcription),
        "translation": "\n".join(all_translation),
    }
